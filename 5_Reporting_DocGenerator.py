
import streamlit as st     
import streamlit.components.v1 as components
from langchain_community.llms import OpenAI
import docx
import datetime
import time
import base64


#BUSINESS CASE (DATA VIZUALIZATION - REPORTING): 
# Ref 1: https://www.structionsite.com/blog/challenges-in-construction-data-management
st.title("A lot of data is wasted")
st.header("Between contractors, sub-contractors, on-site staff, suppliers, and other stakeholders, construction projects generate a massive amount of data. This data can be used by business leaders to gain insight into the successes and failures of their current operations. However, as much as 80% of this data is usually left unstructured and therefore, considered to be “dark data.")
st.image('reporting.png', caption='https://www.structionsite.com/blog/challenges-in-construction-data-management')



#MICROSOFT SOLUTION (REPORT VIZUALIZATION WITH POWER BI):


path_to_html = "/Users/juanrivera/Desktop/chatbot1/Streamlit/pages/11_power_BI.html" 

# Read file and keep in variable
with open(path_to_html,'r') as f: 
    html_data = f.read()

## Show in webpage
components.html(html_data,height=2000)



#MICROSOFT SOLUTION (QUICK DOC GENERATOR WITH CHATGPT):


st.markdown("<h1 style='text-align:justified;font-family:Georgia'>Construction Chatbot - Doc Generator</h1>",unsafe_allow_html=True)

with st.sidebar:
    openai_api_key = st.secrets["auth_token"]
    st.markdown("-------")
    Company_name=st.text_input("What is the name of the company?")
    start_up_description=st.text_input("Please describe whats your statement and the objectives of your report")
    sector = st.multiselect('In which sector of construction the project is?', ["Non - Residential", "Residential"])
    st.markdown("-------")

    generatebutt=st.button("Generate my Report")

    
def generate_report(Company_name, report_date):
    
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading("Report", 0)
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Created For: {Company_name}')
    doc.add_heading(f'Balance of {Company_name} for {sector} sector')
    doc.save('Construction Report.docx')
    data = open('Construction Report.docx', "rb").read()
    encoded = base64.b64encode(data)
    decoded = base64.b64decode(encoded)

    st.download_button('Download Here', decoded, "Construction Report.docx")

def generate_response(input_text):
    llm = OpenAI(temperature=0.3, openai_api_key=openai_api_key)
    output=llm(input_text)
    return output


# User input field
# Initialize chat history


if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({"role": "", "content": "Hey there, I'm OptiLabor Bot , here to help you to create your report \n\
                                      Please put the man inputs of yur report on the left and then ask me about the main purpose of the report"})
# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        with st.spinner('Starting Bot ...'):
            st.markdown(message["content"])

if generatebutt:
    if openai_api_key.startswith('sk-'):
        date_today=datetime.date.today()
        Funding_generated_summary=generate_response(f"I'm currently in the process of exploring funding options for my startup named {Company_name}, and I'd like to gather as much information as possible.The description of my start up is {start_up_description}.\n\
            I'm particularly interested in understanding the various funding sources available to early-stage startups like mine and any specific tips or considerations.I've selected the following  options. \n\
                                    To begin,I'd appreciate an overview of the different types of funding sources that are accessible to my startups in and related to {sector}.Moreover, I'd like to understand the eligibility requirements and criteria that startups typically need to meet for each funding source I've selected. This information will be invaluable as I evaluate which options align with \n\
                                    my startup's current stage and objectives.Preparing strong applications or pitches is crucial when seeking funding. Therefore, I would welcome any advice or tips on how to present a compelling case to potential investors or funding organizations.\n\
                                        Understanding what investors look for can significantly enhance my chances of securing the necessary funds. Networking is often a vital aspect of the fundraising process.\n\
                                        If you could provide strategies for building connections with potential investors or organizations that provide funding, I would greatly appreciate it. Insights into effective networking can be a game-changer.\n\
                                        Additionally, I'd like to be aware of common challenges or pitfalls that startups frequently encounter during the fundraising process.Learning from these experiences can help me avoid potential setbacks and navigate the process more effectively.\n\
                                        Lastly, timing and planning are critical considerations. Insights into when it's the right time to seek funding and how to plan for a successful fundraising campaign would be highly valuable. \n\
                                        If you could also share any relevant resources, articles, or additional advice on this topic, it would be greatly appreciated. Your assistance in this matter is of utmost importance to me as I embark on this funding journey for my startup. Put it in point form and complete each point and up to date specified information.")
        
        Legal_generated_summary=generate_response(f"I am seeking your legal expertise to guide me through the process of launching my startup called {Company_name} in a specific. I would appreciate comprehensive advice that covers all relevant legal requirements, regulations, and considerations unique to this jurisdiction and related to {sector}.\n\
                                    Please provide as much information as possible to ensure a successful and compliant startup launch in this country. Your insights are invaluable in navigating the legal landscape effectively. Put it in point form and complete each point and up to date specified information.")
        
        generate_report(Company_name, date_today)
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')
      
if (prompt := st.chat_input("What is up?")): 
    if  openai_api_key.startswith('sk-'):
        # Display user message in chat message container
        #user=
        with st.chat_message(""):
            st.markdown(prompt)
        # Add user message to chat history
        st.session_state.messages.append({"role": "", "content": prompt})

        # Display assistant response in chat message container
        #assistant=
    # Display assistant response in chat message container
        with st.chat_message(""):
            message_placeholder = st.empty()
            full_response = ""
            with st.spinner('Wait for it...'):
                assistant_response = generate_response(prompt)

        for chunk in assistant_response.split():
            full_response += chunk + " "
            time.sleep(0.05)
            # Add a blinking cursor to simulate typing
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

        # Add assistant response to chat history
        st.session_state.messages.append({"role": "", "content": full_response})
    
    else:
        st.warning('Please enter your OpenAI API key!', icon='⚠')


